import os
import uuid
import json
import io
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template, send_from_directory
try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    psycopg2 = None
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont, ImageOps

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
FORM_SAVE_PATH = os.path.join(BASE_DIR, 'saved_form.json')
HISTORY_PATH   = os.path.join(BASE_DIR, 'saved_history.json')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB max upload

ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ─── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_photo():
    if 'file' not in request.files:
        return jsonify({'error': 'Geen bestand'}), 400
    f = request.files['file']
    if f.filename == '' or not allowed_file(f.filename):
        return jsonify({'error': 'Ongeldig bestand'}), 400
    ext = f.filename.rsplit('.', 1)[1].lower()
    filename = f'{uuid.uuid4().hex}.{ext}'
    f.save(os.path.join(UPLOAD_FOLDER, filename))
    return jsonify({'filename': filename, 'url': f'/uploads/{filename}'})

@app.route('/uploads/<filename>')
def serve_upload(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

# ─── Generate schema PNG from grid data ───────────────────────────────────────

@app.route('/generate-schema', methods=['POST'])
def generate_schema():
    data    = request.get_json()
    floors  = data.get('floors', ['KD', 'BG'])   # bottom-to-top
    columns = int(data.get('columns', 2))
    cells   = data.get('cells', {})

    CELL    = 80
    LABEL_W = 80

    # Rows top-to-bottom: empty header, named floors reversed, then _extra (unlabeled)
    display_rows = [''] + list(reversed(floors)) + ['_extra']
    total_cols   = columns + 1  # +1 always-empty extra col

    img_w = LABEL_W + total_cols * CELL
    img_h = len(display_rows) * CELL

    img  = Image.new('RGB', (img_w, img_h), 'white')
    draw = ImageDraw.Draw(img)

    def load_font(bold=False, size=16):
        candidates = [
            '/System/Library/Fonts/Helvetica.ttc',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf' if bold
                else '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf' if bold
                else '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        ]
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except (OSError, IOError):
                continue
        return ImageFont.load_default()

    font      = load_font(bold=False, size=16)
    font_bold = load_font(bold=True,  size=16)

    BORDER = '#aaaaaa'
    GRAY   = '#d8d8d8'

    for row_idx, floor_label in enumerate(display_rows):
        y = row_idx * CELL

        # Label cell
        is_named  = floor_label and floor_label != '_extra'
        is_extra_row = floor_label == '_extra'
        if is_named:
            draw.rectangle([0, y, LABEL_W - 1, y + CELL - 1], fill=GRAY, outline=BORDER)
            bbox = draw.textbbox((0, 0), floor_label, font=font_bold)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            draw.text((LABEL_W - 10 - tw, y + (CELL - th) // 2), floor_label, font=font_bold, fill='#333333')
        elif is_extra_row:
            # unlabeled row: gray cell, no text
            draw.rectangle([0, y, LABEL_W - 1, y + CELL - 1], fill=GRAY, outline=BORDER)
        else:
            draw.rectangle([0, y, LABEL_W - 1, y + CELL - 1], fill='white', outline=BORDER)

        # Data + extra cells
        for col_idx in range(total_cols):
            x = LABEL_W + col_idx * CELL
            draw.rectangle([x, y, x + CELL - 1, y + CELL - 1], fill='white', outline=BORDER)

            if (is_named or is_extra_row) and col_idx < columns:
                value = cells.get(f'{floor_label}-{col_idx}', '')
                if value:
                    bbox = draw.textbbox((0, 0), value, font=font)
                    tw = bbox[2] - bbox[0]
                    th = bbox[3] - bbox[1]
                    draw.text((x + (CELL - tw) // 2, y + (CELL - th) // 2), value, font=font, fill='black')

    filename = f'{uuid.uuid4().hex}.png'
    img.save(os.path.join(UPLOAD_FOLDER, filename))
    return jsonify({'filename': filename})

@app.route('/save-form', methods=['POST'])
def save_form():
    data = request.get_json()
    with open(FORM_SAVE_PATH, 'w', encoding='utf-8') as fh:
        json.dump(data, fh, ensure_ascii=False, indent=2)
    return jsonify({'ok': True})

@app.route('/load-form')
def load_form():
    if not os.path.exists(FORM_SAVE_PATH):
        return jsonify({})
    with open(FORM_SAVE_PATH, encoding='utf-8') as fh:
        return jsonify(json.load(fh))

# ─── History backend (Postgres when DATABASE_URL set, else JSON file) ──────────

DATABASE_URL = os.environ.get('DATABASE_URL')

def _use_db():
    return bool(DATABASE_URL and psycopg2)

def _db_conn():
    url = DATABASE_URL
    # Railway sometimes gives postgres:// but psycopg2 needs postgresql://
    if url.startswith('postgres://'):
        url = 'postgresql://' + url[len('postgres://'):]
    return psycopg2.connect(url)

def _ensure_table():
    """Create the history table if it doesn't exist (called once at startup)."""
    if not _use_db():
        return
    try:
        with _db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS history (
                        id      TEXT PRIMARY KEY,
                        naam    TEXT,
                        datum   TEXT,
                        formdata TEXT
                    )
                """)
            conn.commit()
    except Exception as e:
        print(f'[history] DB table init failed: {e}')

_ensure_table()

# ── JSON-file fallbacks (local dev) ───────────────────────────────────────────

def _load_history_file():
    if not os.path.exists(HISTORY_PATH):
        return []
    with open(HISTORY_PATH, encoding='utf-8') as fh:
        return json.load(fh)

def _save_history_file(entries):
    with open(HISTORY_PATH, 'w', encoding='utf-8') as fh:
        json.dump(entries, fh, ensure_ascii=False, indent=2)

# ── Public history API ─────────────────────────────────────────────────────────

def _history_insert(entry_id, naam, datum, formdata_json):
    if _use_db():
        with _db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO history (id, naam, datum, formdata) VALUES (%s, %s, %s, %s)",
                    (entry_id, naam, datum, formdata_json)
                )
            conn.commit()
    else:
        entries = _load_history_file()
        entries.append({
            'id': entry_id, 'naam': naam, 'datum': datum,
            'formdata': json.loads(formdata_json),
        })
        _save_history_file(entries)

@app.route('/history')
def get_history():
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute("SELECT id, naam, datum FROM history ORDER BY datum DESC")
                    return jsonify(cur.fetchall())
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify([{k: v for k, v in e.items() if k != 'formdata'}
                        for e in reversed(_load_history_file())])

@app.route('/history/<entry_id>')
def get_history_entry(entry_id):
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("SELECT id, naam, datum, formdata FROM history WHERE id = %s", (entry_id,))
                    row = cur.fetchone()
            if not row:
                return jsonify({}), 404
            return jsonify({
                'id': row[0], 'naam': row[1], 'datum': row[2],
                'formdata': json.loads(row[3]),
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        for e in _load_history_file():
            if e['id'] == entry_id:
                return jsonify(e)
        return jsonify({}), 404

@app.route('/history/<entry_id>', methods=['DELETE'])
def delete_history_entry(entry_id):
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM history WHERE id = %s", (entry_id,))
                conn.commit()
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        _save_history_file([e for e in _load_history_file() if e['id'] != entry_id])
    return jsonify({'ok': True})

# ─── Documents (explicitly saved form data, editable) ─────────────────────────

DOCUMENTS_SAVE_PATH = os.path.join(BASE_DIR, 'saved_documents.json')

def _ensure_documents_table():
    if not _use_db():
        return
    try:
        with _db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS documents (
                        id               TEXT PRIMARY KEY,
                        adres            TEXT,
                        datum_aangemaakt TEXT,
                        datum_aangepast  TEXT,
                        data             TEXT
                    )
                """)
            conn.commit()
    except Exception as e:
        print(f'[documents] DB table init failed: {e}')

_ensure_documents_table()

def _load_documents_file():
    if not os.path.exists(DOCUMENTS_SAVE_PATH):
        return []
    with open(DOCUMENTS_SAVE_PATH, encoding='utf-8') as fh:
        return json.load(fh)

def _save_documents_file(docs):
    with open(DOCUMENTS_SAVE_PATH, 'w', encoding='utf-8') as fh:
        json.dump(docs, fh, ensure_ascii=False, indent=2)

@app.route('/documents')
def list_documents():
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        "SELECT id, adres, datum_aangemaakt, datum_aangepast "
                        "FROM documents ORDER BY datum_aangepast DESC"
                    )
                    return jsonify(cur.fetchall())
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify([
            {k: v for k, v in d.items() if k != 'data'}
            for d in reversed(_load_documents_file())
        ])

@app.route('/documents/<doc_id>')
def get_document(doc_id):
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT id, adres, datum_aangemaakt, datum_aangepast, data "
                        "FROM documents WHERE id = %s", (doc_id,)
                    )
                    row = cur.fetchone()
            if not row:
                return jsonify({}), 404
            return jsonify({
                'id': row[0], 'adres': row[1],
                'datum_aangemaakt': row[2], 'datum_aangepast': row[3],
                'data': json.loads(row[4]),
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        for d in _load_documents_file():
            if d['id'] == doc_id:
                return jsonify(d)
        return jsonify({}), 404

@app.route('/documents', methods=['POST'])
def create_document():
    body     = request.get_json()
    adres    = (body.get('adres') or '').strip() or datetime.now().strftime('%Y-%m-%d %H:%M')
    now      = datetime.now().strftime('%Y-%m-%d %H:%M')
    doc_id   = str(uuid.uuid4())
    data_str = json.dumps(body.get('data', {}), ensure_ascii=False)
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO documents (id, adres, datum_aangemaakt, datum_aangepast, data) "
                        "VALUES (%s, %s, %s, %s, %s)",
                        (doc_id, adres, now, now, data_str)
                    )
                conn.commit()
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        docs = _load_documents_file()
        docs.append({
            'id': doc_id, 'adres': adres,
            'datum_aangemaakt': now, 'datum_aangepast': now,
            'data': body.get('data', {}),
        })
        _save_documents_file(docs)
    return jsonify({'id': doc_id, 'adres': adres, 'datum_aangemaakt': now, 'datum_aangepast': now})

@app.route('/documents/<doc_id>', methods=['PUT'])
def update_document(doc_id):
    body     = request.get_json()
    adres    = (body.get('adres') or '').strip() or datetime.now().strftime('%Y-%m-%d %H:%M')
    now      = datetime.now().strftime('%Y-%m-%d %H:%M')
    data_str = json.dumps(body.get('data', {}), ensure_ascii=False)
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE documents SET adres=%s, datum_aangepast=%s, data=%s WHERE id=%s",
                        (adres, now, data_str, doc_id)
                    )
                conn.commit()
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        docs = _load_documents_file()
        for d in docs:
            if d['id'] == doc_id:
                d['adres'] = adres
                d['datum_aangepast'] = now
                d['data'] = body.get('data', {})
                break
        _save_documents_file(docs)
    return jsonify({'id': doc_id, 'adres': adres, 'datum_aangepast': now})

@app.route('/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    if _use_db():
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM documents WHERE id = %s", (doc_id,))
                conn.commit()
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        _save_documents_file([d for d in _load_documents_file() if d['id'] != doc_id])
    return jsonify({'ok': True})

@app.route('/generate', methods=['POST'])
def generate():
    data = request.get_json()
    buf = io.BytesIO()
    try:
        doc = build_document(data)
        doc.save(buf)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    # Save to history (non-fatal)
    try:
        naam = (data.get('cover', {}).get('straat_huisnrs', '') or '').strip()
        if not naam:
            naam = datetime.now().strftime('%Y-%m-%d %H:%M')
        _history_insert(
            entry_id=str(uuid.uuid4()),
            naam=naam,
            datum=datetime.now().strftime('%Y-%m-%d %H:%M'),
            formdata_json=json.dumps(data, ensure_ascii=False),
        )
    except Exception:
        pass
    buf.seek(0)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='Plaatsingsdocument.docx'
    )

# ─── Document helpers ──────────────────────────────────────────────────────────

def styled_run(para, text, bold=False, italic=False, size=10, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    return run

def set_cell_shading(cell, fill='BFBFBF'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_table_borders(table, border_val='single', sz='4'):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), border_val)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)

def _compress_image(image_path):
    """Resize to max 1920px longest side and return a JPEG BytesIO."""
    MAX = 1920
    with Image.open(image_path) as img:
        img = ImageOps.exif_transpose(img)
        img = img.convert('RGB')
        w, h = img.size
        if max(w, h) > MAX:
            if w >= h:
                img = img.resize((MAX, round(h * MAX / w)), Image.LANCZOS)
            else:
                img = img.resize((round(w * MAX / h), MAX), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=75, optimize=True)
        buf.seek(0)
        return buf

def add_centered_image(doc, image_path, width_cm, max_height_cm=None):
    if not image_path or not os.path.exists(image_path):
        return

    static_dir = os.path.join(BASE_DIR, 'static')
    if image_path.startswith(static_dir):
        img_src = image_path          # static asset — use path directly
    else:
        img_src = _compress_image(image_path)   # user upload — compress

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run()
    if max_height_cm:
        try:
            if isinstance(img_src, io.BytesIO):
                img_src.seek(0)
                with Image.open(img_src) as pil_img:
                    w_px, h_px = pil_img.size
                img_src.seek(0)
            else:
                with Image.open(img_src) as pil_img:
                    w_px, h_px = pil_img.size
            aspect = w_px / h_px
            h_at_max_w = width_cm / aspect
            if h_at_max_w <= max_height_cm:
                run.add_picture(img_src, width=Cm(width_cm))
            else:
                if isinstance(img_src, io.BytesIO):
                    img_src.seek(0)
                run.add_picture(img_src, height=Cm(max_height_cm))
        except Exception:
            if isinstance(img_src, io.BytesIO):
                img_src.seek(0)
            run.add_picture(img_src, width=Cm(width_cm))
    else:
        run.add_picture(img_src, width=Cm(width_cm))

def photo_path(filename):
    if not filename:
        return None
    p = os.path.join(UPLOAD_FOLDER, filename)
    return p if os.path.exists(p) else None

def add_section_heading(doc, number, title, size=13):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    styled_run(para, f'{number}: {title}', bold=True, size=size)

def add_para(doc, text, bold=False, italic=False, size=10, indent_cm=None, space_after=2):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    styled_run(para, text, bold=bold, italic=italic, size=size)
    return para

def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

def cell_para(cell, text='', bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if text:
        styled_run(para, text, bold=bold, italic=italic, size=size)
    return para

def add_gray_row(table, label, col_span=None):
    """Add a full-width gray header row with a label."""
    row = table.add_row()
    cell = row.cells[0]
    # merge all cells in the row
    for c in row.cells[1:]:
        cell = cell.merge(c)
    cell.text = ''
    set_cell_shading(cell, 'BFBFBF')
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    styled_run(para, label, bold=True, size=9)

def add_field_row(table, label, value, col_widths=None):
    """Add a label-value row (2 cells)."""
    row = table.add_row()
    cell_para(row.cells[0], label, bold=False, size=9)
    cell_para(row.cells[1], value or '', size=9)

# ─── Cover Page ────────────────────────────────────────────────────────────────

def build_cover(doc, data):
    cover = data.get('cover', {})

    # Photo
    img = photo_path(cover.get('foto_vooraanzicht'))
    if img:
        add_centered_image(doc, img, 8)
    else:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        styled_run(para, '[Vooraanzicht foto]', italic=True, size=9)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(4)
    styled_run(title_para, 'Plaatsingsdocument', bold=True, size=24)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after = Pt(12)
    styled_run(sub_para, 'Vooraanzichtpand', bold=False, size=16)

    # Info table – 2 columns
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    usable = 15.92
    set_col_widths(table, [6.0, usable - 6.0])

    def add_row(label, value):
        row = table.add_row()
        cell_para(row.cells[0], label, bold=True, size=10)
        cell_para(row.cells[1], str(value) if value else '', size=10)

    add_row('Installatiecode:', 'Plaatsingsdocument---')
    add_row('HB nr:', cover.get('hb_nr', ''))
    add_row('DP:', cover.get('dp', ''))
    add_row('Projectnr:', '14246')
    add_row('Straat + huisnrs:', cover.get('straat_huisnrs', ''))
    add_row('BAG Pand ID:', cover.get('bag_pand_id', ''))
    add_row('Postcode + Plaats:', cover.get('postcode_plaats', ''))
    add_row('AP-gebied:', cover.get('ap_gebied', ''))
    add_row('DP-gebied:', cover.get('dp_gebied', ''))
    add_row('Datum schouw:', cover.get('datum_schouw', ''))
    add_row('Aannemer:', cover.get('aannemer', ''))
    add_row('Doc. opgesteld door:', 'B Sayilir')
    add_row('Naam schouwer:', 'Bayram')

    add_page_break(doc)

# ─── Inhoudsopgave ─────────────────────────────────────────────────────────────

def build_toc(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)
    styled_run(para, 'Inhoudsopgave', bold=True, size=16)

    items = [
        'Locatie en installatiegegevens',
        'Locatie',
        'Beslisboom',
        'Aanzicht complex eenheden',
        'Beschrijving tracé en kabelroute',
        'Schematische tekening',
        'Beschrijving aanleg',
        'Materiaalstaat',
        'Overeenkomst/instemmingsbesluit',
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        styled_run(p, f'{i}.  {item}', size=11)

    add_page_break(doc)

# ─── Sectie 1: Locatie en installatiegegevens ──────────────────────────────────

def build_s1(doc, data):
    s1 = data.get('s1', {})
    add_section_heading(doc, 1, 'Locatie en installatiegegevens')

    # Main table – 2 columns (label | value)
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    set_col_widths(table, [6.5, 9.42])

    def gray(label):
        add_gray_row(table, label)

    def row(label, value, fixed=False):
        r = table.add_row()
        cell_para(r.cells[0], label, size=9)
        cell_para(r.cells[1], value or '', size=9, bold=fixed)

    # ── Algemeen ──
    gray('Algemeen')
    row('Datum site survey:', s1.get('datum_site_survey', ''))
    row('Glasvezelvoorbereider:', s1.get('glasvezelvoorbereider', ''))
    row('Tel. nr:', s1.get('tel_nr', ''))
    row('Type Netwerkconcept:', '')
    row('Type pand:', 'App. complex', fixed=True)
    row('Type woningen (Huur/koop):', '')
    row('Aantal aansluitingen:', s1.get('aantal_aansluitingen', ''))
    row('Aantal verdiepingen:', s1.get('aantal_verdiepingen', ''))

    # ── Horizontal etage table (13 columns) ──
    # First add a label paragraph (gray-styled heading in the flow)
    lbl_p = doc.add_paragraph()
    lbl_p.paragraph_format.space_before = Pt(0)
    lbl_p.paragraph_format.space_after = Pt(0)
    lbl_tbl = doc.add_table(rows=1, cols=1)
    lbl_cell = lbl_tbl.rows[0].cells[0]
    lbl_cell.width = Cm(15.92)
    set_cell_shading(lbl_cell, 'BFBFBF')
    set_table_borders(lbl_tbl)
    lbl_cell_para = lbl_cell.paragraphs[0]
    lbl_cell_para.paragraph_format.space_before = Pt(1)
    lbl_cell_para.paragraph_format.space_after = Pt(1)
    styled_run(lbl_cell_para, 'Aantal woningen per etage', bold=True, size=9)

    etage_cols      = ['KD', 'BG', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11']
    etage_data_keys = ['KD', 'BG', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11']
    etage_vals = s1.get('woningen_per_etage', {})

    # 14 cols: 1 empty label col + 13 etage cols
    etage_tbl = doc.add_table(rows=2, cols=14)
    set_table_borders(etage_tbl)
    label_w = 0.5
    data_w  = round((15.92 - label_w) / 13, 3)
    set_col_widths(etage_tbl, [label_w] + [data_w] * 13)

    # Header row: col 0 = empty gray; cols 1-13 = KD, BG, 1-11
    set_cell_shading(etage_tbl.rows[0].cells[0], 'BFBFBF')
    for i, col in enumerate(etage_cols):
        c = etage_tbl.rows[0].cells[i + 1]
        set_cell_shading(c, 'BFBFBF')
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(1)
        cp.paragraph_format.space_after  = Pt(1)
        styled_run(cp, col, bold=True, size=8)

    # Data row: col 0 = empty; cols 1-13 = values
    data_row = etage_tbl.rows[1]
    for i, key in enumerate(etage_data_keys):
        cp = data_row.cells[i + 1].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(1)
        cp.paragraph_format.space_after  = Pt(1)
        styled_run(cp, etage_vals.get(key, ''), size=9)

    # Resume main 2-column table for remaining general fields
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    set_col_widths(table, [6.5, 9.42])

    def gray(label):
        add_gray_row(table, label)

    def row(label, value, fixed=False):
        r = table.add_row()
        cell_para(r.cells[0], label, size=9)
        cell_para(r.cells[1], value or '', size=9, bold=fixed)

    row('Geschatte hoogte van de site:', s1.get('geschatte_hoogte', ''))
    row('Hoogwerker nodig?:', 'Nee', fixed=True)
    row('Type hoogwerker:', '')
    row('Voorwerk mogelijk?:', 'Nee', fixed=True)
    row('Type voorbereiding:', 'FTU', fixed=True)
    row('Bewoners afhankelijk van elkaar?:', s1.get('bewoners_afhankelijk', ''))
    row('Asbest:', 'NEE (In te vullen door gebouw eigenaar/beheerder) / bij NEE geen verder onderzoek nodig / bij JA zie volgende regel', fixed=True)
    row('Brandscheiding tekeningen:', 'JA/NEE/WEET NIET (In te vullen door gebouw eigenaar/beheerder)', fixed=True)
    row('Brandvertragende afdichtingen:', 'NEE (In te vullen door gebouw eigenaar/beheerder)', fixed=True)
    row('Toelichting / Afwijking:', '')

    # ── Contactinformatie ──
    gray('Contactinformatie')
    row('VVE/WOCO/VGE:', s1.get('contact_vve_naam', ''))
    row('Contactpersoon beheer:', s1.get('contact_persoon_beheer', ''))
    row('Functie:', s1.get('functie', ''))
    row('Telefoon:', s1.get('telefoon', ''))
    row('Mobiel:', s1.get('mobiel', ''))
    row('Emailadres:', s1.get('emailadres', ''))
    row('Meerdere eigenaren:', s1.get('meerdere_eigenaren', ''))
    row('Gevestigd te:', s1.get('gevestigd_te', ''))
    row('Toegang site via:', s1.get('toegang_site_via', ''))

    # ── Aanwezigen tijdens survey ──
    gray('Aanwezigen tijdens survey')
    aanwezigen = s1.get('aanwezigen', [])
    if aanwezigen:
        for a in aanwezigen:
            row(a.get('naam', ''), a.get('telefoon', ''))
    else:
        row('Naam:', '')
        row('Telefoonnummer:', '')

    # ── Installatiegegevens ──
    gray('Installatiegegevens')
    row('Locatie in- of opvoerpunt:', s1.get('locatie_invoerpunt', ''))
    row('Aantal invoergaten nodig:', s1.get('aantal_invoergaten', ''))
    row('Stijgpunt:', s1.get('stijgpunt', ''))
    row('Mantelbuis of koof toegankelijk?:', s1.get('mantelbuis_toegankelijk', ''))
    row('Bestaande doorvoer gebruiken?:', s1.get('bestaande_doorvoer', ''))
    row('Locatie glasvezelaansluitpunt:', s1.get('locatie_gvap', ''))
    row('Locatie CAI AOP:', s1.get('locatie_cai_aop', ''))
    row('Locatie ISRA:', s1.get('locatie_isra', ''))
    row('Locatie meterkast:', s1.get('locatie_meterkast', ''))
    row('230V aanwezig?:', s1.get('volt_aanwezig', ''))
    row('Locatie hoofdafsluiters:', s1.get('locatie_hoofdafsluiters', ''))
    row('Locatie Modem:', s1.get('locatie_modem', ''))
    row('Locatie FTU:', s1.get('locatie_ftu', ''))
    row('Brandwerende afdichting aanwezig:', s1.get('brandwerende_afdichting', ''))

    add_page_break(doc)

# ─── Sectie 2: Locatie ─────────────────────────────────────────────────────────

def build_s2(doc, data):
    s2 = data.get('s2', {})
    add_section_heading(doc, 2, 'Locatie')

    add_para(doc, 'Locatie complex:', bold=True, size=10)
    add_para(doc, 'Locatie bovenaanzicht invoerpunten:', bold=True, size=10)

    img = photo_path(s2.get('foto_kaart'))
    add_centered_image(doc, img, 12, max_height_cm=8) if img else None

    tekst = s2.get('tekst_locatie', '')
    if tekst:
        add_para(doc, tekst, size=10)

    add_page_break(doc)

# ─── Sectie 3: Beslisboom ──────────────────────────────────────────────────────

def build_s3(doc, data):
    s3 = data.get('s3', {})
    add_section_heading(doc, 3, 'Beslisboom')

    # Always use the fixed beslisboom image from static/
    beslisboom_path = os.path.join(BASE_DIR, 'static', 'beslisboom.png')
    if os.path.exists(beslisboom_path):
        add_centered_image(doc, beslisboom_path, 10)
    else:
        # Gray placeholder if file not yet saved
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        cell.width = Cm(14)
        set_cell_shading(cell, 'BFBFBF')
        tr = cell._tc.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(4 * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(30)
        styled_run(para, 'Sla beslisboom.png op in static/ map', italic=True, size=10)

    uitleg = s3.get('uitleg_beslisboom', '')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    styled_run(p, 'Uitleg keuze beslisboom door aannemer:', bold=True, size=10)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    styled_run(p2, uitleg or '', size=10)

    add_page_break(doc)

# ─── Sectie 4: Aanzicht complex eenheden ──────────────────────────────────────

def build_s4(doc, data):
    s4 = data.get('s4', {})
    add_section_heading(doc, 4, 'Aanzicht complex eenheden')

    fotos = s4.get('fotos', [])
    # backwards-compat: single foto_aanzicht key
    if not fotos and s4.get('foto_aanzicht'):
        fotos = [{'foto': s4['foto_aanzicht'], 'caption': s4.get('caption', '')}]

    for item in fotos:
        img = photo_path(item.get('foto'))
        if img:
            add_centered_image(doc, img, 12, max_height_cm=8)
        caption = item.get('caption', '')
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            styled_run(p, caption, italic=True, size=9)

    add_page_break(doc)

# ─── Sectie 5: Beschrijving tracé en kabelroute ────────────────────────────────

def build_s5(doc, data):
    s5 = data.get('s5', {})
    add_section_heading(doc, 5, 'Beschrijving tracé en kabelroute')

    def add_photo_subsection(title, items):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        styled_run(p, title, bold=True, italic=True, size=11)

        for item in items:
            img = photo_path(item.get('foto'))
            add_centered_image(doc, img, 12, max_height_cm=8) if img else None
            tekst = item.get('tekst', '')
            if tekst:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(2)
                p2.paragraph_format.space_after = Pt(8)
                styled_run(p2, tekst, size=10)

    add_photo_subsection('Tracé routes (buiten):', s5.get('buiten', []))
    add_photo_subsection('Route inpandig:', s5.get('inpandig', []))
    add_photo_subsection('FTU locatie voorbeeld:', s5.get('ftu', []))

    p_ftu = doc.add_paragraph()
    p_ftu.paragraph_format.space_before = Pt(4)
    p_ftu.paragraph_format.space_after  = Pt(6)
    styled_run(p_ftu, 'Let op: aannemer en glasvezelpartij zijn niet verantwoordelijk voor het realiseren van een 230 volt voorziening.', italic=True, size=10)

    add_page_break(doc)

# ─── Sectie 6: Schematische tekening ──────────────────────────────────────────

def build_s6(doc, data):
    s6 = data.get('s6', {})
    add_section_heading(doc, 6, 'Schematische tekening')

    img = photo_path(s6.get('schema_image'))
    if img:
        add_centered_image(doc, img, 14, max_height_cm=20)

    add_page_break(doc)

# ─── Sectie 7: Beschrijving aanleg (volledig vast) ────────────────────────────

def build_s7(doc, data):
    s7 = data.get('s7', {})
    add_section_heading(doc, 7, 'Beschrijving aanleg')

    # Planning
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'Planning:', bold=True, italic=True, size=10)
    add_para(doc, (
        'De woningen worden per streng gepland (verticale stijging). Dit betekent dat wij, m.b.t. de doorvoer, '
        'van alle bewoners afhankelijk zijn om de woningen vanaf de begane grond tot aan de hoogste etage '
        'gelijktijdig aan te kunnen sluiten. Bewoners worden doormiddel van een afspraakbrief op de hoogte '
        'gesteld, op welke datum en dagdeel de aansluiting wordt gerealiseerd.'
    ), size=10, space_after=6)

    # Afwerking
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'Afwerking van kabels en invoer/doorvoer gaten.', bold=True, italic=True, size=10)

    add_para(doc, (
        'De invoeren die gemaakt worden voor het glasvezeltraject in de woning/ het complex, worden na het '
        'invoeren van de glasvezelkabel zorgvuldig dichtgemaakt met daarvoor gebruikelijke Stoppac.'
    ), size=10, space_after=3)
    add_para(doc, (
        'Dit afdichtingsmateriaal is waterdicht en gas-belemmerend en voldoet aan de Nederlandse NEN-norm 2768.'
    ), size=10, space_after=3)
    add_para(doc, (
        'Alle door ons te maken doorvoeren door brandwerende muren en plafonds (scheidingsconstructies) worden '
        'volgens de normen brandwerend afgewerkt.'
    ), size=10, space_after=3)
    add_para(doc, (
        'Met betrekking tot de doorvoeren en vanuit onze zorgplicht, artikel 1.16 Naar Bouwbesluit en/of de '
        'Woningwet, maken wij u erop attent dat de eisen aan de weerstand tegen branddoorslag, brandoverslag en '
        'rookdoorgang zoals deze gesteld worden in de afdelingen/artikelen voorschrijven dat bestaande en nieuwe '
        'doorvoeren (al dan niet na het aanbrengen en/of wijzigingen) inclusief samenhangende voorzieningen '
        '(zoals brandwerende brandmachetten) adequaat beheerd, onderhouden en gecontroleerd dienen te worden. '
        'Mocht blijken dat niet meer aan de voorschriften van het bouwbesluit wordt voldaan, dienen deze '
        'tekortkomingen direct worden hersteld.'
    ), size=10, space_after=6)

    # Montagevoorschrift
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'Montagevoorschrift;', bold=True, italic=True, size=10)

    montage_items = [
        'Bestrating en/of groen worden opengelegd bij invoerpunt(-en).',
        'Doorvoer bij elk pand wordt gasdicht achtergelaten.',
        'Alle kabels worden voorzien van het juiste huisnummer en vezelnummer.',
        'Alle kabels worden afgedicht d.m.v. einddop.',
        'Alle geboorde doorvoeren worden gasdicht achtergelaten.',
        'Geopende bestrating/plantsoenen worden dichtgelegd en waar nodig hersteld.',
        'Na afloop werkzaamheden wordt de werkplek schoon en opgeruimd achtergelaten.',
    ]
    for i, item in enumerate(montage_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        styled_run(p, f'{i}. {item}', size=10)

    # Werkomschrijving
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'Werkomschrijving;', bold=True, italic=True, size=10)

    werk_items = [
        'Vanaf tracé naar de gevel van het pand graven/pneumatisch boren.',
        'Bij de invoerpunten worden de kabels via de stijgingen naar de locatie opgevoerd.',
        'De kabels worden afgewerkt op een BOP. De Pre-connectorized kabels stijgen via een nieuw aan te brengen rvs kabelgoot en of overzetgoot en worden per verdieping naar binnen geboord.',
        'De kabels komen hier tegen de gevel op rol. De kabels worden doorgevoerd door de bestaande of nieuwe doorvoer. De kabels stijgen d.m.v. een nieuw aan te brengen rvs goot en of overzetgoot en worden per verdieping naar binnen geboord.',
        'Bij de opvoerpunten worden de kabels middels RVS goot naar de woonkamers of meterkast opgevoerd',
        'Glasvezelaansluitpunten (FTU\'s) worden volgens de omschreven locatie in de Beslisboom geplaatst.',
        'Routering wordt volgens Omschrijving tracé en kabelroute aangelegd.',
    ]
    for item in werk_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run_check = p.add_run('✓ ')
        run_check.font.name = 'Segoe UI Symbol'
        run_check.font.size = Pt(10)
        styled_run(p, item, size=10)

    # Strenglijsten
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'Strenglijsten', bold=True, size=12)

    add_para(doc, 'Geschatte in- of opvoerpunt per woning', size=10, space_after=4)

    # Strenglijsten table – 7 columns
    strengen = s7.get('strenglijsten', [])
    tbl = doc.add_table(rows=1 + max(len(strengen), 5), cols=7)
    set_table_borders(tbl)
    # Column widths: Adres(4) Hnr(1.2) Lengte(1.5) Invoer(2.5) St.p(1.5) DP(1.5) Streng ID(3.72) = 15.92
    set_col_widths(tbl, [4.0, 1.2, 1.5, 2.5, 1.5, 1.5, 3.72])

    headers = ['Adres', 'Hnr', 'Lengte', 'Invoer', 'St.p', 'DP', 'Streng ID']
    data_keys = ['adres', 'hnr', 'lengte', 'invoer', 'stp', 'dp', 'streng_id']

    for i, hdr in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_shading(c, 'BFBFBF')
        cell_para(c, hdr, bold=True, size=9)

    for ri in range(1, 1 + max(len(strengen), 5)):
        row_data = strengen[ri - 1] if ri - 1 < len(strengen) else {}
        cells = tbl.rows[ri].cells
        for ci, key in enumerate(data_keys):
            cell_para(cells[ci], row_data.get(key, ''), size=9)

    add_para(doc, (
        'De geschatte in- of opvoerpunt van de glasvezelkabel is de lengte die op rol gelegd bij het invoerpunt '
        'van het complex. Let op: aantal meters is inclusief overlengte.'
    ), size=10, space_after=6)

    # Buis en gootvoorzieningen
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    styled_run(p, 'In- en uitpandige buis- en gootvoorzieningen', bold=True, italic=True, size=10)

    add_para(doc, (
        'Met de ondertekening van dit installatiedocument verleen ik, als eigenaar van de in- en uitpandige '
        'buis- en gootvoorzieningen t.b.v. de KPN en CAI bekabeling, toestemming om gebruik hiervan te mogen '
        'maken. Dit om de omschreven glasvezelbekabeling in het betreffende complex aan te brengen. Tevens '
        'verleen ik toestemming voor het gebruik maken voor het plaatsen van hoogwerkers/auto\'s en/of '
        'materieel op (eigen) grond om de werkzaamheden uit te kunnen voeren.'
    ), size=10, space_after=6)

    add_page_break(doc)

# ─── Sectie 8: Materiaalstaat ─────────────────────────────────────────────────

def build_s8(doc, data):
    add_section_heading(doc, 8, 'Materiaalstaat')

    tbl = doc.add_table(rows=1, cols=4)
    set_table_borders(tbl)
    set_col_widths(tbl, [5.5, 2.5, 4.5, 3.42])

    headers = ['Omschrijving', 'Aantal', 'Type', 'Huidige te gebruiken?']
    for i, hdr in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_shading(c, 'BFBFBF')
        cell_para(c, hdr, bold=True, size=10)

    # 5 empty rows
    for _ in range(5):
        tbl.add_row()

    add_para(doc, (
        'Standaard verbruiksmaterialen; Pluggen, schroeven, trekontlasting, kit, Stopaq, krimpkous, '
        'flexibele buis, tube, pur schuim, beugels en vulbusjes.'
    ), size=10, space_after=6)

    add_page_break(doc)

# ─── Sectie 9: Overeenstemming ────────────────────────────────────────────────

def build_s9(doc, data):
    s1 = data.get('s1', {})
    vve_naam = s1.get('contact_persoon_beheer', '') or s1.get('contact_vve_naam', '')

    add_section_heading(doc, 9, 'Overeenstemming / Ondertekening')

    add_para(doc, (
        'Onderstaande partijen gaan akkoord met realisatie van het glasvezelnetwerk conform de uitvoering en '
        'gebruik making van de materialen zoals omschreven in dit document*.'
    ), size=10, space_after=12)

    # 2-column signature table
    sig_tbl = doc.add_table(rows=5, cols=2)
    remove_table_borders(sig_tbl)
    set_col_widths(sig_tbl, [7.96, 7.96])

    def sig_cell(row_idx, col_idx, text, bold=False, italic=False, size=10):
        cell = sig_tbl.rows[row_idx].cells[col_idx]
        cell_para(cell, text, bold=bold, italic=italic, size=size)

    # Row 0: signature lines
    sig_cell(0, 0, '________________________________')
    sig_cell(0, 1, '________________________________')

    # Row 1: role labels
    entity_label = s1.get('contact_vve_naam', '').strip() or 'VvE'
    sig_cell(1, 0, 'Aannemer', bold=True, size=10)
    sig_cell(1, 1, entity_label, bold=True, size=10)

    # Row 2: "Naam ondergetekende:"
    sig_cell(2, 0, 'Naam ondergetekende:', size=10)
    sig_cell(2, 1, 'Naam ondergetekende:', size=10)

    # Row 3: name value
    sig_cell(3, 0, '')
    sig_cell(3, 1, vve_naam, size=10)

    # Row 4: Datum
    sig_cell(4, 0, 'Datum:', size=10)
    sig_cell(4, 1, 'Datum:', size=10)

    # Footnote
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        '* Indien er om uitvoeringstechnische redenen afgeweken dient te worden van de uitvoeringswijze zoals '
        'beschreven in het getekende installatiedocument, zal in overleg met de contactpersoon ter plaatse een '
        'nieuw technisch uitvoerbaar tracé worden bepaald.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.italic = True

# ─── Main document builder ────────────────────────────────────────────────────

def build_document(data):
    doc = Document()

    # Global defaults
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.0)

    # Footer on every page
    section.footer_distance = Cm(1.0)
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)
    r = footer_para.add_run('Plaatsingsdocument---Versie 2.2')
    r.font.name = 'Calibri'
    r.font.size = Pt(7)
    r.bold = True

    # Build each section
    build_cover(doc, data)
    build_toc(doc)
    build_s1(doc, data)
    build_s2(doc, data)
    build_s3(doc, data)
    build_s4(doc, data)
    build_s5(doc, data)
    build_s6(doc, data)
    build_s7(doc, data)
    build_s8(doc, data)
    build_s9(doc, data)

    return doc

# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
